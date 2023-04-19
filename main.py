from PyPDF2 import PdfReader
import openai
import datetime
from docx import Document
from docx.shared import Pt, Cm
import requests
from bs4 import BeautifulSoup

import utils

openai.api_key = utils.open_file('openai_api_key.txt')

def extract_text_from_pdf(pdf)-> str:
    """Extract the text from a pdf file and return a string"""
    reader = PdfReader(pdf)
    page = reader.pages[0]
    return page.extract_text().strip().replace('É', 'E')

def extract_text_from_url(url:str)-> str:
    """Scraps a website and returns the text as a string"""
    page = requests.get(url)
    soup = BeautifulSoup(page.content, 'html.parser')
    return ' '.join([el.text for el in soup.find_all('p') if 'cookie' not in el.text])


def is_a_cv(cv:str)-> bool:
    cv_discussion = [{"role": "system", "content": "You are a CV classifier model. You are given a text as an input, you ouput 1 if it looks like a CV, 0 if not. Only 1 or 0, nothing else."}]
    experience, _ = utils.gpt3_chat(message=cv, messages=cv_discussion)
    return '1' in experience


def is_a_job_offer(offer:str)-> bool:
    offer_discussion = [{"role": "system", "content": "You are a job offer classifier model. You are given a text as an input, you ouput 1 if it looks like a job description, 0 if not. Only 1 or 0, nothing else."}]
    experience, _ = utils.gpt3_chat(message=offer, messages=offer_discussion)
    return '1' in experience

def extract_content_from_CV(cv:str)-> dict:
    """Extracts specific content blocks from the CV and output them in a dictionary"""
    # Experience 
    cv_discussion = [
        {"role": "system", "content": "You are a smart CV parser, you extract specific elements from curriculum vitae and output them in a specific format. When asked, output the desired format only, no human response."},
        {"role": "user", "content": f"Here is the CV that you are going to parse: <<{cv}>>"},
        {"role": "assistant", "content": "What specific elements do you want me to extract?"},
        {"role": "user", "content": f"I want you to extract the professional experiences and the tasks/missions performed at each position"},
        {"role": "assistant", "content": "What is the expected output format?"},
    ]
    experience_prompt = "Expected output format:\n[Company name]: [Job title]\n- Duration: [Start date] - [End date]\n- Tasks:\n  - [Task description 1]\n  - [Task description 2]\n  - [Task description 3]\n...\n\nAnswer the formatted output only, add nothing else, be exhaustive and copy paste the missions descriptions as they are. Freelance experiences count (skip the [Company Name] part). Do it now."
    experience, cv_discussion = utils.gpt3_chat(message=experience_prompt, messages=cv_discussion)
    # Education
    education_prompt = "Perfect. Do the same with the education part now. Expected output format:\n[School Name]:\n- Duration: [Start date] - [End date]\n- Degree: [Degree]\n- [Courses]: [List courses]\n...\n\nAnswer the formatted output only, add nothing else."
    education, cv_discussion = utils.gpt3_chat(message=education_prompt, messages=cv_discussion)
    # Lastest degree and school
    degree_and_school_names_prompt = "What is the lastest degree and the last school of the candidate? Anwser in this format '[Degree] at [School Name]'. If it ended before 2023, add 'Not graduating' at the end of your answer"
    degree_and_school_names, cv_discussion = utils.gpt3_chat(message=degree_and_school_names_prompt, messages=cv_discussion)
    is_graduating = 'not graduating' not in degree_and_school_names.lower()
    # Current position
    current_position_prompt = "What is the first job title and company name writen on this CV? Answer only '[Job title] at [Company name]'. If it is a freelance activity, just answer 'Freelance [Job title]'. If this job ended before 2023, add 'Not working' at the end of your answer"
    current_position, cv_discussion = utils.gpt3_chat(message=current_position_prompt, messages=cv_discussion)
    is_working = 'not working' not in current_position.lower()
    # Name
    name_prompt = "Perfect. What is the name on the CV? Just answer the name, nothing else. For example if the name is 'Bill Gates', just output 'Bill Gates'"
    name, cv_discussion = utils.gpt3_chat(message=name_prompt, messages=cv_discussion)
    name = utils.clean_word_output(name)
    return {'name': name,
            'experience': experience,
            'education': education,
            'degree_and_school_names': degree_and_school_names,
            'is_graduating': is_graduating,
            'current_position': current_position,
            'is_working': is_working}


def extract_content_from_offer(offer:str)-> dict:
    """Extracts specific content blocks from the job offer and output them in a dictionary"""
    # Missions
    offer_discussion = [
        {"role": "system", "content": "You are a smart Job Description Parser. You extract specific elements from job offers and output them in a specific format. You answer in English only, and translate if needed."},
        {"role": "user", "content": f"Here is the job description that you are going to parse: <<{offer}>>"},
        {"role": "assistant", "content": "What specific elements do you want me to extract?"}
    ]
    missions_prompt = f"Extract, sum up and list the 3 main missions of this job. Answer in English only and go straight to the 3 points, no introduction sentence."
    missions, offer_discussion = utils.gpt3_chat(message=missions_prompt, messages=offer_discussion)
    # Job title
    job_title_prompt = f"Extract the job title (if the job title is not clearly mentioned, imagine what it could be). Output the job title only, do not make a sentence. For example, if the job title is Marketing Lead, you must output 'Marketing Lead' only."
    job_title, offer_discussion = utils.gpt3_chat(message=job_title_prompt, messages=offer_discussion)
    job_title = utils.clean_word_output(job_title)
    # Company name
    company_name_prompt = f"Extract the name of the company. If the company name is not clearly mentioned, anwser 'XXX'. Output the company name only, do not make a sentence. For example, if the company name is Microsoft, you must output 'Microsoft' only."
    company_name, offer_discussion = utils.gpt3_chat(message=company_name_prompt, messages=offer_discussion)
    company_name = utils.clean_word_output(company_name)
    # Motivations
    motivations_prompt = f"Extract, sum up and list the 3 main reasons why it is an amazing opportunity for the candidate to work as a {job_title} at {company_name}. Insist on the company's strengths (ex: leader in its domain, great corporate culture, fast-growing environment) and on its mission. Answer in English only and go straight to the 3 points, no introduction sentence."
    motivations, offer_discussion = utils.gpt3_chat(message=motivations_prompt, messages=offer_discussion)
    return {'company_name': company_name,
            'job_title': job_title,
            'motivations': motivations,
            'missions': missions}
    
def create_prompt_filling_dict(cv_content:dict, offer_content:dict)-> dict:  
    """Merges cv_content and offer_content into one dictionary used to fill the prompt"""
    return {'<<EXPERIENCE>>': cv_content['experience'],
            '<<EDUCATION>>': cv_content['education'],
            '<<CANDIDATE_NAME>>': cv_content['name'],
            '<<DEGREE_AND_SCHOOL_NAMES>>': cv_content['degree_and_school_names'],
            '<<CURRENT_POSITION>>': cv_content['current_position'],
            '<<IS_GRADUATING>>': cv_content['is_graduating'],
            '<<IS_WORKING>>': cv_content['is_working'],
            '<<MISSIONS>>': offer_content['missions'],
            '<<JOB_TITLE>>': offer_content['job_title'],
            '<<COMPANY_NAME>>': offer_content['company_name'],
            '<<MOTIVATIONS>>': offer_content['motivations']}


def create_cover_letter_blocks(prompt_filling_dict:dict)-> dict:
    # General information
    candidate_name = prompt_filling_dict['<<CANDIDATE_NAME>>']
    job_title = prompt_filling_dict['<<JOB_TITLE>>']
    company_name = prompt_filling_dict['<<COMPANY_NAME>>']
    # Date
    now = datetime.datetime.now()
    date_block = f"{now.strftime('%B')} {now.strftime('%d')}, {now.strftime('%Y')}"
    # Subject
    prompt_path = "prompts/cl_subject.txt"
    subject_block = utils.fill_prompt(utils.open_file(prompt_path), prompt_filling_dict)
    # Salutations
    salutations_block = 'Dear Sir or Madam,'
    # Intro 
    is_graduating, is_working = prompt_filling_dict['<<IS_GRADUATING>>'], prompt_filling_dict['<<IS_WORKING>>']
    if is_graduating and is_working:
        prompt_path = "prompts/cl_intro_graduate_working.txt"
    elif is_graduating:
        prompt_path = "prompts/cl_intro_graduate.txt"
    elif is_working:
        prompt_path = "prompts/cl_intro_working.txt"
    else:
        prompt_path = "prompts/cl_intro_standard.txt"
    prompt = utils.fill_prompt(utils.open_file(prompt_path), prompt_filling_dict)
    intro_block = utils.gpt3_completion(prompt, temp=0, tokens=100)
    # Motivations
    prompt_path = "prompts/cl_motivations.txt"
    prompt = utils.fill_prompt(utils.open_file(prompt_path), prompt_filling_dict)
    motivations_block = utils.gpt3_completion(prompt, temp=0, tokens=300).replace('\n', '')
    # Experiences
    prompt_path = "prompts/cl_experiences.txt"
    experience_transition = 'I am confident that my previous experiences demonstrate my alignment with these three motivations and explain my eagerness to work as <<JOB_TITLE>> at <<COMPANY_NAME>>. '
    experience_transition = utils.fill_prompt(experience_transition, prompt_filling_dict)
    prompt = utils.fill_prompt(utils.open_file(prompt_path), prompt_filling_dict)
    experiences_block = utils.gpt3_completion(prompt, temp=0, tokens=300).replace('\n', '')
    experiences_block = experience_transition + experiences_block
    # Education
    prompt_path = "prompts/cl_education.txt"
    prompt = utils.fill_prompt(utils.open_file(prompt_path), prompt_filling_dict)
    education_block = utils.gpt3_completion(prompt, temp=0, tokens=300).replace('\n', '')
    # Missions
    prompt_path = "prompts/cl_missions.txt"
    prompt = utils.fill_prompt(utils.open_file(prompt_path), prompt_filling_dict)
    missions_block = utils.gpt3_completion(prompt, temp=0, tokens=300).replace('\n', '')
    # Greeting
    prompt_path = "prompts/cl_greeting.txt"
    greeting_block = utils.fill_prompt(utils.open_file(prompt_path), prompt_filling_dict)
    # Closing
    closing_block = 'Yours sincerely,'
    return {'candidate_name': candidate_name,
            'job_title': job_title,
            'company_name': company_name,
            'date_block': date_block,
            'subject_block': subject_block,
            'salutations_block': salutations_block,
            'intro_block': intro_block, 
            'motivations_block': motivations_block,
            'experiences_block': experiences_block,
            'education_block': education_block,
            'missions_block': missions_block,
            'greeting_block': greeting_block,
            'closing_block': closing_block}


def create_cover_letter(cl_blocks: dict, export_file=False)-> tuple:
    """Creates the cover letter and outputs it as a word file"""
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    date = document.add_paragraph(cl_blocks['date_block'])
    date.alignment = 2
    document.add_paragraph()

    subject = document.add_paragraph()
    subject.add_run(cl_blocks['subject_block']).bold = True
    subject.alignment = 3

    salutations = document.add_paragraph(cl_blocks['salutations_block'])

    for block in [cl_blocks['intro_block'], cl_blocks['motivations_block'], cl_blocks['experiences_block'], cl_blocks['education_block'], cl_blocks['missions_block'], cl_blocks['greeting_block']]:
        p = document.add_paragraph(block)
        p.alignment = 3

    closing = document.add_paragraph(cl_blocks['closing_block'])

    document.add_paragraph()
    signature = document.add_paragraph(cl_blocks['candidate_name'])
    signature.alignment = 2

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(9.5)

    job_title = cl_blocks['job_title'].replace(' ', '_').lower()
    company_name = cl_blocks['company_name'].replace(' ', '_').lower()
    now = datetime.datetime.now()
    current_datetime = now.strftime('%Y%m%d') + '_' + now.strftime('%H%M%S')
    document_name = f"cl_{company_name}_{job_title}_{current_datetime}"
    if export_file:
        document_path = f'results/{document_name}.docx'
        document.save(document_path)
    return document, document_name

if __name__ == '__main__':
    url = input('Enter the URL:')
    print('scraping the url and extracting content...')
    offer = extract_text_from_url(url)
    print('extracting content from CV...')
    cv = extract_text_from_pdf('files/cv.pdf')
    print('checking that the file provided is a CV and the url refers a job offer...')
    if not is_a_cv(cv):
        print('Sorry buddy, but the CV you uploaded does not look like a CV 🤯')
    elif not is_a_job_offer(offer):
        print('Sorry buddy, but the URL you provided does not look like job offer 🤯')
    else:
        print('parsing the CV to extract specific content...')
        cv_content = extract_content_from_CV(cv)
        print('parsing the job offer to extract specific content...')
        offer_content = extract_content_from_offer(offer)
        prompt_filling_dict = create_prompt_filling_dict(cv_content, offer_content)
        print('creating the different cover letter blocks...')
        cl_blocks = create_cover_letter_blocks(prompt_filling_dict)
        print('creating the word document...')
        document, document_name = create_cover_letter(cl_blocks)
